# # -*- coding: utf-8 -*-
# # @Time         : 2023/7/16
# # @Author       : Yuyang Zhan
# # @File         : aav20.py
# # @Software     : PyCharm
# # @E-mail       : 02yonaa@gmail.com

# #8.16更改比对方式为blast比对，并添加一个txt文件来输出与多个ref拥有相同最高匹配率的序列
# #9.10添加process_word_file函数，区分aav和质粒
# #10.4添加利用xml信息判断正反向序列，如果是反向则将其正过来
# #10.5添加预警信息，确保所有提供的reference都转变成了fasta格式，并添加一个计算时间的函数，计算从第一条序列匹配第一个reference开始一直到全部匹配完需要的时间
# #10.14质粒文件搜索首先搜索新的标识"StartOfSeq"和"EndOfSeq"，如果未找到，再搜索旧的"begin"和"end"
# #10.26添加fastq和docx预警，并且将输出到same.txt中的序列也同时输出到same.fastq中；以及添加delete_xml_files函数，在程序完成后删除生成的大量xml文件
# #2024.7.18修改计算匹配分数的分母，以及nomatch在填入的时候才打开
# #2024.7.29修改计算匹配得分的方式，由原先的最大匹配得分/seq，改为计算覆盖率：覆盖长度/序列长度
# import os
# import collections
# from docx import Document
# import subprocess
# from Bio.Blast import NCBIXML
# from Bio.Seq import Seq
# import re
# import time
# import sys
# from Bio import SeqIO
# import tempfile

# # 目标序列
# ITR_L = 'CCTGCAGGCAGCTGCGCGCT'
# ITR_R = 'AGCGCGCAGCTGCCTGCAGG'

# input_folder = r"./"

# fastq_files = [os.path.join(input_folder, filename) for filename in os.listdir(input_folder) if
#                filename.endswith(".fastq")]
# if not fastq_files:
#     print("no .fastq file")

# docx_files = [os.path.join(input_folder, filename) for filename in os.listdir(input_folder) if
#               filename.endswith(".docx")]
# if not docx_files:
#     print("no .docx file")

# seq_no_match_file_path = os.path.join(input_folder, "undemultiplexed.fastq")
# same_txt_file_path = os.path.join(input_folder, "same.txt")
# seq_no_match_file = open(seq_no_match_file_path, 'w')
# same_txt_file = open(same_txt_file_path, 'w')

# #aav
# '''
# def process_word_file(file_path):
#     doc = Document(file_path)
#     text = ""
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             for char in run.text:
#                 if char.isalpha():
#                     text += char.upper()
#     start_index = text.find(ITR_L)
#     end_index = text.find(ITR_R)
#     if start_index != -1 and end_index != -1:
#         return text[start_index:end_index + len(ITR_R)]
#     else:
#         return None
# '''

# #质粒
# def process_word_file(file_path):
#     doc = Document(file_path)
#     text = ""
#     found_seq_start = False
#     found_seq_end = False

#     for paragraph in doc.paragraphs:
#         if "startofseq" in paragraph.text.lower():
#             found_seq_start = True
#             continue
#         if "endofseq" in paragraph.text.lower() and found_seq_start:
#             found_seq_end = True
#             break
#         if found_seq_start:
#             for run in paragraph.runs:
#                 for char in run.text:
#                     if char.isalpha():
#                         text += char.upper()


#     if not found_seq_start or not found_seq_end:
#         # 如果没有找到"StartOfSeq"和"EndOfSeq"，则寻找"begin"和"end"
#         found_begin = False
#         for paragraph in doc.paragraphs:
#             if "begin" in paragraph.text.lower():
#                 found_begin = True
#                 continue
#             if "end" in paragraph.text.lower() and found_begin:
#                 break
#             if found_begin:
#                 for run in paragraph.runs:
#                     for char in run.text:
#                         if char.isalpha():
#                             text += char.upper()

#     if text:
#         return text
#     else:
#         return None


# def merge_intervals(intervals):
#     sorted_intervals = sorted(intervals, key=lambda x: x[0])
#     merged_intervals = []
#     for interval in sorted_intervals:
#         if not merged_intervals or merged_intervals[-1][1] < interval[0]:
#             merged_intervals.append(interval)
#         else:
#             merged_intervals[-1][1] = max(merged_intervals[-1][1], interval[1])
#     return merged_intervals

# def calculate_coverage(seq, ref_seq_path, seq_header):
#     # 改版：使用临时文件，避免 header 太长产生超长文件名
#     tmp = tempfile.NamedTemporaryFile(mode='w', suffix=".xml", delete=False)
#     output_path = tmp.name
#     tmp.close()

#     cmd = ['blastn', '-query', '-', '-subject', ref_seq_path, '-outfmt', '5']
#     with open(output_path, 'w') as output_file:
#         subprocess.run(cmd, input=seq, text=True, stdout=output_file)

#     intervals = []
#     reverse_match = False

#     with open(output_path, "r") as result_handle:
#         blast_record = NCBIXML.read(result_handle)
#         for alignment in blast_record.alignments:
#             for hsp in alignment.hsps:
#                 start, end = (hsp.query_start, hsp.query_end) if hsp.query_start <= hsp.query_end else (hsp.query_end, hsp.query_start)
#                 intervals.append([start, end])

#     merged_intervals = merge_intervals(intervals)
#     covered_length = sum(interval[1] - interval[0] + 1 for interval in merged_intervals)
#     seq_length = blast_record.query_length
#     coverage_percentage = (covered_length / seq_length) * 100 if seq_length > 0 else 0.0

#     print(f"Sequence Length: {seq_length}")
#     print(f"Covered Length: {covered_length}")
#     print(f"Coverage Percentage: {coverage_percentage:.2f}%")
#     os.remove(output_path)
#     return coverage_percentage, reverse_match

# def write_fastq(output_file, seq_header, seq):
#     output_file.write(seq_header + '\n')
#     output_file.write(seq + '\n')
#     output_file.write('+\n')
#     output_file.write('I' * len(seq) + '\n')


# def check_fasta_generated(docx_files, ref_fasta_files):
#     generated_fasta_names = [os.path.splitext(os.path.basename(fasta_file))[0] for fasta_file in ref_fasta_files]

#     for docx_file in docx_files:
#         expected_name = os.path.splitext(os.path.basename(docx_file))[0]

#         if expected_name not in generated_fasta_names:
#             print(f"No fasta sequence generated for Word file: {docx_file}")
#             return False

#     return True


# # Generate fasta files and store paths in ref_fasta_files
# ref_fasta_files = []
# for docx_file in docx_files:
#     file_identifier = os.path.splitext(os.path.basename(docx_file))[0]
#     ref_text = process_word_file(docx_file)
#     if ref_text:
#         ref_seq_path = os.path.join(input_folder, f"{file_identifier}.fasta")
#         ref_fasta_files.append(ref_seq_path)
#         with open(ref_seq_path, "w") as ref_fasta_file:
#             ref_fasta_file.write(">Reference\n")
#             ref_fasta_file.write(ref_text)

# # Check if for every docx file a fasta file is generated, if not, stop the code execution.
# if not check_fasta_generated(docx_files, ref_fasta_files):
#     sys.exit("Not all fasta sequences were generated. Exiting...")


# def process_fastq_file(fastq_file, output_folder, ref_fasta_files):
#     # 打开nomatch文件进行写入
#     seq_no_match_file_path = os.path.join(output_folder, "nomatch.fastq")

#     with open(fastq_file, 'r') as input_file:
#         lines = input_file.readlines()
#         for i in range(0, len(lines), 4):
#             seq_header = lines[i].strip()
#             seq = lines[i + 1].strip()

#             best_match_file = ""
#             max_match_rate = 0.0
#             matched_refs = []

#             # 对每一个fasta参考文件进行比对
#             for ref_seq_path in ref_fasta_files:
#                 file_identifier = os.path.splitext(os.path.basename(ref_seq_path))[0]
#                 match_rate, reverse_match = calculate_coverage(seq, ref_seq_path, seq_header)
#                 # 如果是反向匹配，则获取反向互补序列
#                 # if reverse_match:
#                 #     seq_obj = Seq(seq)
#                 #     seq = str(seq_obj.reverse_complement())

#                 # 更新最高匹配分数和对应的参考文件
#                 if match_rate > max_match_rate:
#                     best_match_file = file_identifier
#                     max_match_rate = match_rate
#                     matched_refs = [file_identifier]
#                 elif match_rate == max_match_rate:
#                     matched_refs.append(file_identifier)

#             # 处理匹配结果
#             if max_match_rate <= 95.0:
#                 # 写入nomatch文件
#                 with open(seq_no_match_file_path, 'a') as seq_no_match_file:
#                     write_fastq(seq_no_match_file, seq_header, seq)  
#             else:
#                 # 选择匹配的fasta文件写入
#                 if len(matched_refs) == 1:
#                     ref_file_identifier = matched_refs[0]
#                     with open(os.path.join(output_folder, f"{ref_file_identifier}.fastq"), 'a') as output_file:
#                         write_fastq(output_file, seq_header, seq)
#                 else:
#                     with open(os.path.join(output_folder, "same.txt"), 'a') as same_txt_file:
#                         same_txt_file.write(f"Seq Name: {seq_header}\n")
#                         same_txt_file.write(f"Sequence: {seq}\n")
#                         same_txt_file.write(f"Score: {max_match_rate}%\n")
#                         same_txt_file.write(f"Matching Refs: {', '.join(matched_refs)}\n\n")
#             delete_xml_files(os.path.dirname(fastq_file))


# def process_all_fastq_files(fastq_files, output_folder, ref_fasta_files):
#     # Here we store the timestamp at which we start processing the first sequence
#     start_time = time.time()

#     for fastq_file in fastq_files:
#         process_fastq_file(fastq_file, output_folder, ref_fasta_files)

#     # Here we store the timestamp at which we finish processing the last sequence
#     end_time = time.time()

#     # The elapsed time is the difference between these two timestamps
#     elapsed_time = end_time - start_time

#     # Finally, we print the elapsed time, converting it to minutes for convenience
#     print(f"Processing took {elapsed_time:.2f} seconds ({elapsed_time / 60:.2f} minutes)")


# def delete_xml_files(directory):
#     for filename in os.listdir(directory):
#         if filename.endswith(".xml"):
#             os.remove(os.path.join(directory, filename))


# output_folder = "./output"
# os.makedirs(output_folder, exist_ok=True)

# # Calling the function that processes all files and times its execution
# process_all_fastq_files(fastq_files, output_folder, ref_fasta_files)

# delete_xml_files(input_folder)

# -*- coding: utf-8 -*-
import os
import collections
from docx import Document
import subprocess
from Bio.Blast import NCBIXML
from Bio.Seq import Seq
import re
import time
import sys
from Bio import SeqIO
import tempfile

# ---------- 仅此处改动 ---------- #
SAMPLE_RE = re.compile(r'[PS]\d{3,5}', re.IGNORECASE)  # 捕获 P**** 或 S****
def extract_sample_id(name: str) -> str:
    """
    从文件名或字符串中提取最后一个形如 P1234 / S5678 的编号。
    若未找到则返回原串。
    """
    hits = SAMPLE_RE.findall(name)
    return hits[-1] if hits else name
# -------------------------------- #
def rename_docx_files(docx_list):
    """
    把 docx 文件重命名为 <样本号>.docx，返回重命名后的路径列表。
    若已符合命名规则则跳过。
    """
    renamed = []
    for path in docx_list:
        dirname, fname = os.path.split(path)
        stem = os.path.splitext(fname)[0]
        sample_id = extract_sample_id(stem)          # P**** / S****
        new_path = os.path.join(dirname, f"{sample_id}.docx")
        if path != new_path:
            # 如目标文件已存在，可在此处理冲突；这里直接覆盖旧的同名文件
            if os.path.exists(new_path):
                os.remove(new_path)
            os.rename(path, new_path)
        renamed.append(new_path)
    return renamed

ITR_L = 'CCTGCAGGCAGCTGCGCGCT'
ITR_R = 'AGCGCGCAGCTGCCTGCAGG'

input_folder = "./"

fastq_files = [os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".fastq")]
if not fastq_files:
    print("no .fastq file")

# docx_files = [os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".docx")]
# if not docx_files:
#     print("no .docx file")
raw_docx_files = [os.path.join(input_folder, f) for f in os.listdir(input_folder) if f.endswith(".docx")]
if not raw_docx_files:
    print("no .docx file")

# ★ 第一步：重命名 Word 文件
docx_files = rename_docx_files(raw_docx_files)

def process_word_file(file_path):
    doc = Document(file_path)
    text = ""
    found_seq_start = False
    found_seq_end = False
    for paragraph in doc.paragraphs:
        if "startofseq" in paragraph.text.lower():
            found_seq_start = True
            continue
        if "endofseq" in paragraph.text.lower() and found_seq_start:
            found_seq_end = True
            break
        if found_seq_start:
            for run in paragraph.runs:
                text += ''.join([c.upper() for c in run.text if c.isalpha()])
    if not text:
        found_begin = False
        for paragraph in doc.paragraphs:
            if "begin" in paragraph.text.lower():
                found_begin = True
                continue
            if "end" in paragraph.text.lower() and found_begin:
                break
            if found_begin:
                for run in paragraph.runs:
                    text += ''.join([c.upper() for c in run.text if c.isalpha()])
    return text if text else None

def merge_intervals(intervals):
    intervals.sort(key=lambda x: x[0])
    merged = []
    for interval in intervals:
        if not merged or merged[-1][1] < interval[0]:
            merged.append(interval)
        else:
            merged[-1][1] = max(merged[-1][1], interval[1])
    return merged

def calculate_coverage(seq, ref_seq_path, seq_header):
    tmp = tempfile.NamedTemporaryFile(mode='w', suffix=".xml", delete=False)
    output_path = tmp.name
    tmp.close()
    cmd = ['blastn', '-query', '-', '-subject', ref_seq_path, '-outfmt', '5']
    with open(output_path, 'w') as output_file:
        subprocess.run(cmd, input=seq, text=True, stdout=output_file)
    intervals = []
    with open(output_path, "r") as result_handle:
        blast_record = NCBIXML.read(result_handle)
        for alignment in blast_record.alignments:
            for hsp in alignment.hsps:
                start, end = sorted([hsp.query_start, hsp.query_end])
                intervals.append([start, end])
    merged = merge_intervals(intervals)
    covered = sum(e - s + 1 for s, e in merged)
    length = blast_record.query_length
    os.remove(output_path)
    return (covered / length * 100 if length > 0 else 0.0), False

def write_fastq(output_file, seq_header, seq):
    output_file.write(seq_header + '\n')
    output_file.write(seq + '\n')
    output_file.write('+\n')
    output_file.write('I' * len(seq) + '\n')

def check_fasta_generated(docx_files, ref_fasta_files):
    generated = [os.path.splitext(os.path.basename(f))[0] for f in ref_fasta_files]
    for docx_file in docx_files:
        expected = extract_sample_id(os.path.splitext(os.path.basename(docx_file))[0])
        if expected not in generated:
            print(f"Missing fasta for: {docx_file}")
            return False
    return True

# 生成参考 fasta
ref_fasta_files = []
for docx_file in docx_files:
    base_name = os.path.splitext(os.path.basename(docx_file))[0]
    sample_id = extract_sample_id(base_name)      
    seq = process_word_file(docx_file)
    if seq:
        # path = os.path.join(input_folder, f"{clean_name}.fasta")
        path = os.path.join(input_folder, f"{sample_id}.fasta")  # <— 仅样本号
        ref_fasta_files.append(path)
        with open(path, "w") as f:
            f.write(">Reference\n")
            f.write(seq)

if not check_fasta_generated(docx_files, ref_fasta_files):
    sys.exit("Not all fasta sequences were generated. Exiting...")

def process_fastq_file(fastq_file, output_folder, ref_fasta_files):
    nomatch_path = os.path.join(output_folder, "nomatch.fastq")
    with open(fastq_file, 'r') as f:
        lines = f.readlines()
    for i in range(0, len(lines), 4):
        header = lines[i].strip()
        seq = lines[i+1].strip()
        best = ""
        max_rate = 0.0
        matches = []
        for ref_path in ref_fasta_files:
            file_id = os.path.splitext(os.path.basename(ref_path))[0]
            rate, _ = calculate_coverage(seq, ref_path, header)
            if rate > max_rate:
                best = file_id
                max_rate = rate
                matches = [file_id]
            elif rate == max_rate:
                matches.append(file_id)
        if max_rate <= 95:
            with open(nomatch_path, 'a') as nf:
                write_fastq(nf, header, seq)
        else:
            if len(matches) == 1:
                outname = extract_sample_id(matches[0]) 
                with open(os.path.join(output_folder, f"{outname}.fastq"), 'a') as of:
                    write_fastq(of, header, seq)
            else:
                same_path = os.path.join(output_folder, "same.txt")
                same_fastq = os.path.join(output_folder, "same.fastq")
                with open(same_path, 'a') as sf, open(same_fastq, 'a') as sfq:
                    sf.write(f"Seq Name: {header}\n")
                    sf.write(f"Sequence: {seq}\n")
                    sf.write(f"Score: {max_rate:.2f}%\n")
                    sf.write(f"Matching Refs: {', '.join([extract_sample_id(m) for m in matches])}\n\n")
                    write_fastq(sfq, header, seq)
        delete_xml_files(os.path.dirname(fastq_file))

def process_all_fastq_files(fastq_files, output_folder, ref_fasta_files):
    start = time.time()
    for f in fastq_files:
        process_fastq_file(f, output_folder, ref_fasta_files)
    print(f"Processing took {time.time() - start:.2f} seconds.")

def delete_xml_files(directory):
    for f in os.listdir(directory):
        if f.endswith(".xml"):
            os.remove(os.path.join(directory, f))

output_folder = "./output"
os.makedirs(output_folder, exist_ok=True)

process_all_fastq_files(fastq_files, output_folder, ref_fasta_files)
delete_xml_files(input_folder)
